from flask import Flask, request, jsonify
import os
import re
import tempfile
import io
from collections import defaultdict

import pdfplumber
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from pymupdf4llm import LlamaMarkdownReader
from docx import Document

app = Flask(__name__)

#####################
# PDF 처리 관련 함수
#####################

def estimate_font_weight(text):
    """
    텍스트의 폰트 굵기를 추정합니다.
    대문자와 공백만 있으면 bold로 간주합니다.
    """
    if re.match(r"^[A-Z\s]+$", text.strip()):
        return "bold"
    return "normal"

def cluster_positions(positions, tolerance=2.0):
    """
    위치 값들을 tolerance 범위 내에서 클러스터링하여 클러스터 중심값 리스트를 반환합니다.
    """
    clusters = []
    for pos in sorted(positions):
        found = False
        for cluster in clusters:
            if abs(cluster[0] - pos) <= tolerance:
                # 평균값 업데이트
                cluster[0] = (cluster[0] * cluster[1] + pos) / (cluster[1] + 1)
                cluster[1] += 1
                found = True
                break
        if not found:
            clusters.append([pos, 1])
    return [cluster[0] for cluster in clusters]

def detect_common_y_positions(documents, page_count, tolerance=2.0, threshold=0.8):
    """
    여러 페이지에서 y 좌표들이 군집화되어 반복된다면 해당 클러스터 중심값을 반환합니다.
    """
    y_positions = []
    page_occurrences = defaultdict(set)
    for doc in documents:
        page = doc.metadata.get("page_number", 0)
        y = doc.metadata.get("bbox", [0, 0, 0, 0])[1]
        y_positions.append(y)
        page_occurrences[y].add(page)
    clusters = cluster_positions(y_positions, tolerance)
    common_y = set()
    for cluster in clusters:
        count = 0
        for y, pages in page_occurrences.items():
            if abs(y - cluster) <= tolerance:
                count += len(pages)
        if count >= page_count * threshold:
            common_y.add(round(cluster, 1))
    return common_y

def detect_repeated_footer_positions(pdf_path, documents, tolerance=2.0, threshold=0.8, bottom_ratio=0.9):
    """
    각 페이지 하단의 텍스트 y 좌표를 군집화하여, 반복되는 footer 영역의 y 좌표 집합을 반환합니다.
    """
    doc = fitz.open(pdf_path)
    page_count = doc.page_count
    y_positions = []
    page_occurrences = defaultdict(set)
    for doc_item in documents:
        page_num = doc_item.metadata.get("page_number", 0)
        page_idx = page_num - 1
        if page_idx < 0 or page_idx >= page_count:
            continue
        page_height = doc[page_idx].rect.height
        y_top = doc_item.metadata.get("bbox", [0, 0, 0, 0])[1]
        if y_top >= page_height * bottom_ratio:
            y_positions.append(y_top)
            page_occurrences[y_top].add(page_num)
    clusters = cluster_positions(y_positions, tolerance)
    common_y = set()
    for cluster in clusters:
        count = 0
        for y, pages in page_occurrences.items():
            if abs(y - cluster) <= tolerance:
                count += len(pages)
        if count >= page_count * threshold:
            common_y.add(round(cluster, 1))
    return common_y

def extract_text_by_page_with_llama(pdf_path):
    """
    LlamaMarkdownReader를 사용하여 PDF의 텍스트와 메타정보를 추출합니다.
    반복되는 Header/Footer 영역은 제거합니다.
    """
    reader = LlamaMarkdownReader()
    documents = reader.load_data(pdf_path)
    page_count = max(doc.metadata.get("page_number", 0) for doc in documents)
    common_y_positions = detect_common_y_positions(documents, page_count)
    repeated_footers = detect_repeated_footer_positions(pdf_path, documents)
    remove_y_positions = common_y_positions.union(repeated_footers)
    tolerance = 2.0

    page_items = defaultdict(list)
    for doc_item in documents:
        page_num = doc_item.metadata.get("page_number", 0)
        bbox = doc_item.metadata.get("bbox", [0, 0, 0, 0])
        y_top = round(bbox[1], 1)
        if any(abs(y_top - pos) < tolerance for pos in remove_y_positions):
            continue

        color = doc_item.metadata.get("font_color", "#000000")
        bgcolor = doc_item.metadata.get("bg_color", "#FFFFFF")
        font_size = bbox[3] - bbox[1]
        link = doc_item.metadata.get("uri", None)
        content = doc_item.text.strip()
        if link:
            content = f"[{content}]({link})"
        item = {
            "type": "text",
            "content": content,
            "x": bbox[0],
            "y": bbox[1],
            "font_size": font_size,
            "font_weight": estimate_font_weight(doc_item.text),
            "font_color": color,
            "bg_color": bgcolor
        }
        page_items[page_num].append(item)
    return page_items

def extract_text_from_pdfplumber_or_ocr(pdf_path):
    """
    pdfplumber로 텍스트를 추출하고, 텍스트가 없으면 pytesseract로 OCR 처리합니다.
    """
    doc = fitz.open(pdf_path)
    results = defaultdict(list)
    with pdfplumber.open(pdf_path) as pdf:
        for page_index, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                results[page_index + 1].append({
                    "type": "text",
                    "content": text.strip(),
                    "x": 0,
                    "y": 0,
                    "font_size": 12,
                    "font_weight": "normal",
                    "font_color": "#000000",
                    "bg_color": "#FFFFFF"
                })
            else:
                pix = doc.load_page(page_index).get_pixmap(dpi=300)
                image = Image.open(io.BytesIO(pix.tobytes()))
                ocr_text = pytesseract.image_to_string(image, lang='eng')
                if ocr_text.strip():
                    results[page_index + 1].append({
                        "type": "text",
                        "content": ocr_text.strip(),
                        "x": 0,
                        "y": 0,
                        "font_size": 12,
                        "font_weight": "normal",
                        "font_color": "#000000",
                        "bg_color": "#FFFFFF"
                    })
    return results

def convert_table_to_markdown(table):
    """
    2차원 배열 형태의 표 데이터를 Markdown 표 형식으로 변환합니다.
    """
    def format_cell(cell):
        return cell.strip() if cell else ""
    if not table or not table[0]:
        return ""
    header = "| " + " | ".join(format_cell(cell) for cell in table[0]) + " |"
    separator = "| " + " | ".join(["---"] * len(table[0])) + " |"
    rows = [
        "| " + " | ".join(format_cell(cell) for cell in row) + " |"
        for row in table[1:] if row
    ]
    return "\n".join([header, separator] + rows)

def classify_heading(item):
    """
    폰트 크기와 bold 여부에 따라 Markdown 헤더 또는 일반 텍스트로 변환합니다.
    """
    size = item.get("font_size", 0)
    bold = (item.get("font_weight", "normal") == "bold")
    content = item['content'].strip()
    font_color = item.get("font_color", "#000000")
    bg_color = item.get("bg_color", "#FFFFFF")
    highlight = (font_color.lower() not in ["#000000", "#000", "black"]) or (bg_color.lower() != "#ffffff")
    style_prefix = "**" if (highlight or bold) else ""
    styled_content = f"{style_prefix}{content}{style_prefix}"
    
    if bold:
        if size >= 24 or (size >= 20):
            return f"# {styled_content}"
        elif size >= 18:
            return f"## {styled_content}"
        else:
            return f"### {styled_content}"
    else:
        if size >= 24:
            return f"# {styled_content}"
        elif size >= 20:
            return f"## {styled_content}"
        elif size >= 14:
            return f"### {styled_content}"
        else:
            return styled_content

def process_text_item(item):
    """
    텍스트 아이템 내에서 bold 상태이고 행 전체가 짧은 경우 소제목(헤더)로 처리합니다.
    """
    content = item["content"].strip()
    if item.get("font_weight", "normal") == "bold":
        lines = content.split("\n")
        header_lines = []
        body_lines = []
        for line in lines:
            stripped_line = line.strip()
            if stripped_line and len(stripped_line) < 50:
                header_lines.append(f"#### {stripped_line}")
            else:
                if stripped_line:
                    body_lines.append(stripped_line)
        if header_lines and body_lines:
            return "\n".join(header_lines) + "\n\n" + "\n".join(body_lines)
        elif header_lines:
            return "\n".join(header_lines)
        else:
            return content
    else:
        paragraphs = re.split(r'\n\s*\n', content)
        if not paragraphs:
            return ""
        first = paragraphs[0].strip()
        if len(first) < 50:
            header = classify_heading({**item, "content": first})
            body = "\n\n".join(p.strip() for p in paragraphs[1:] if p.strip())
            if body:
                return header + "\n\n" + body
            else:
                return header
        else:
            return "\n\n".join(p.strip() for p in paragraphs if p.strip())

def normalize_string(s):
    """
    문자열의 불필요한 공백을 제거하여 정규화합니다.
    """
    return " ".join(s.split())

def extract_combined_markdown(pdf_path):
    """
    PDF 파일을 처리하여,
      1) LlamaMarkdownReader로 텍스트 추출 (헤더/푸터 제거)
      2) pdfplumber로 표 및 OCR 텍스트 추출
      3) 표 영역에 포함된 텍스트 및 중복되는 표 헤더 제거
      4) (y, x) 좌표 기준 원본 순서를 재현한 Markdown을 생성
    """
    llama_text_items = extract_text_by_page_with_llama(pdf_path)
    fallback_text_items = extract_text_from_pdfplumber_or_ocr(pdf_path)
    footer_pattern = re.compile(r"^Page\s*\d+(\s*(of|/)\s*\d+)?\s*$", re.IGNORECASE)
    markdown_parts = []
    margin = 5
    
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_elements = []
            if llama_text_items.get(page_num):
                text_items = llama_text_items[page_num]
            elif fallback_text_items.get(page_num):
                text_items = fallback_text_items[page_num]
            else:
                text_items = []
            
            table_boxes = []
            table_headers = []
            tables = page.find_tables()
            for table in tables:
                if not table:
                    continue
                bbox = table.bbox
                table_boxes.append(bbox)
                extracted_table = table.extract()
                md_table = convert_table_to_markdown(extracted_table)
                if extracted_table and len(extracted_table) > 0:
                    header_row = extracted_table[0]
                    header_text = " ".join(cell.strip() for cell in header_row if cell).strip()
                    table_headers.append(header_text)
                page_elements.append({
                    "type": "table",
                    "content": "\n" + md_table,
                    "x": bbox[0],
                    "y": bbox[1]
                })
            
            filtered_text_items = []
            for item in text_items:
                x, y = item.get("x", 0), item.get("y", 0)
                inside_table = any(
                    bbox[0] - margin <= x <= bbox[2] + margin and
                    bbox[1] - margin <= y <= bbox[3] + margin
                    for bbox in table_boxes
                )
                if inside_table:
                    continue
                content = item.get("content", "").strip()
                normalized_content = normalize_string(content)
                if any(normalize_string(header) == normalized_content for header in table_headers):
                    continue
                filtered_text_items.append(item)
            
            for item in filtered_text_items:
                if footer_pattern.match(item.get("content", "").strip()):
                    continue
                page_elements.append(item)
            
            sorted_elements = sorted(page_elements, key=lambda e: (e["y"], e["x"]))
            page_section = []
            for element in sorted_elements:
                if element["type"] == "text":
                    page_section.append(process_text_item(element))
                elif element["type"] == "table":
                    page_section.append(element["content"])
            if page_section:
                markdown_parts.extend(page_section)
                markdown_parts.append("\n---\n")
    return "\n".join(markdown_parts)

#####################
# Word 파일 처리 함수
#####################

def convert_docx_to_markdown(docx_path):
    """
    python-docx를 사용하여 Word(.docx) 파일의 단락과 표를 Markdown 형식으로 변환합니다.
    단락의 스타일(Heading)과 bold 여부를 참고하여 제목(헤더) 처리를 합니다.
    """
    document = Document(docx_path)
    markdown_lines = []
    
    # 단락 처리
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 스타일이 Heading이면 Markdown 제목으로 변환
        if para.style.name.startswith("Heading"):
            level = para.style.name.split()[-1] if para.style.name.split()[-1].isdigit() else 1
            markdown_lines.append("#" * int(level) + " " + text)
        else:
            # bold 상태이고 길이가 짧은 경우 소제목으로 처리
            is_bold = any(run.bold for run in para.runs if run.text.strip())
            if is_bold and len(text) < 50:
                markdown_lines.append("#### " + text)
            else:
                markdown_lines.append(text)
        markdown_lines.append("")  # 빈 줄 추가

    # 표 처리
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        md_table = convert_table_to_markdown(rows)
        markdown_lines.append(md_table)
        markdown_lines.append("")
    
    return "\n".join(markdown_lines)

#####################
# Flask API 엔드포인트
#####################

@app.route("/convert", methods=["POST"])
def convert_file_to_markdown():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    filename = file.filename.lower()
    ext = os.path.splitext(filename)[1]
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        if ext in [".pdf"]:
            markdown = extract_combined_markdown(tmp_path)
        elif ext in [".doc", ".docx"]:
            markdown = convert_docx_to_markdown(tmp_path)
        else:
            return jsonify({"error": "Unsupported file type"}), 400

        return jsonify({"markdown": markdown})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8021, debug=True)
